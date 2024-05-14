from ast import keyword
import time
from turtle import up
import docx
from time import sleep
import webbrowser
from flask import current_app
import requests
import statistics
from httpx import request
import pyttsx3
import qrcode
import cv2
import wolframalpha
import datetime
from winsound import PlaySound
import os
from datetime import datetime
from pyautogui import click
from keyboard import _physically_pressed_keys
from keyboard import press
from keyboard import write
from keyboard import press_and_release
import speech_recognition as sr
from geopy.distance import great_circle
from geopy.geocoders import Nominatim
import geocoder
from notifypy import Notify
from notifypy import notify

engine = pyttsx3.init('sapi5')
voices = engine.getProperty('voices')
#print(voices[0].id)
engine.setProperty('voice', voices[0].id)
engine.setProperty('rate', 180-200)

def speak(audio):
    engine.say(audio)
    engine.runAndWait()

def takeCommand():
    #It takes microphone input from the user and returns string output

    r = sr.Recognizer()
    with sr.Microphone() as source:
        print("Listenning ...")
        r.pause_threshold = 1
        audio = r.listen(source)
    
    try:
        print("Rcognizing...")
        query = r.recognize_google(audio, language='en-in')
        print(f"User said: {query}\n")

    except Exception as e:
        print(e)
        print("Say that again please...")
        return "None"
    return query

def mean(l1):
    print("Mean for given data --> ", statistics.mean(l1))
    speak(f"Mean for given data is {statistics.mean(l1)}")

def median(l1):
    print("Median for given data --> ", statistics.median(l1))
    speak(f"Median for given data is {statistics.median(l1)}")

def mode(l1):
    print("Mode for given data --> ", statistics.mode(l1))
    speak(f"Mode for given data is {statistics.mode(l1)}")

def variance(l1):
    print("Variance for given data --> ", statistics.variance(l1))
    speak(f"Variance for given data is {statistics.variance(l1)}")

def stdev(l1):
    print("Standard Deviation for given data --> ", statistics.stdev(l1))
    speak(f"Standard Deviation for given data is {statistics.stdev(l1)}")

def range(l1):
    Maximum = max(l1)
    Minimum = min(l1)
    Range = Maximum-Minimum    
    print(f"Range of {l1} is", Range)

def QR_Code_Generator(s,img_name):  
    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.ERROR_CORRECT_L,
        box_size=10,
        border=4,
    )
    qr.add_data(s)
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white")
    img.save(img_name)

def QR_reader(nm):
    img = cv2.imread(nm)
    det = cv2.QRCodeDetector()
    hid_data = det.detectAndDecode(img)
    print(hid_data)

def WolfRam(query):
    api_key = "PUGTVE-UGAUTTV3t2"
    requester = wolframalpha.Client(api_key)
    requested = requester.query(query)
    try:
        Answer = next(requested.results).text
        return Answer
    except:
        print("No data found! Try again")
        speak("No data found! Try again")

def Temp(query):
    Term = str(query)
    Term = Term.replace("jarvis", "")
    Term = Term.replace("in", "")
    Term = Term.replace("what is the", "")
    Term = Term.replace("temperature", "")
    temp_query = str(Term)
    if 'outside' in temp_query:
        var1 = "Temperature in Delhi"
        answer = WolfRam(var1)
        speak(f"{var1} is {answer}.")
    else:
        var2 = "Temperature in " + temp_query
        answ = WolfRam(var2)
        speak(f"{var2} in {answ}.")

def Alarm(query):
    extracted_time = open('C:\\Users\\HIMI\\Desktop\\Jarvis\\data.txt','rt')
    time = extracted_time.read()
    Time = str(time)

    delete_time = open('C:\\Users\\HIMI\\Desktop\\Jarvis\\data.txt','r+')
    delete_time.truncate(0)
    delete_time.close()

    def RingerNow(time):
        time_to_set = str(time)
        time_now = time_to_set.replace("jarvis","")
        time_now = time_to_set.replace("set alarm for","")
        time_now = time_to_set.replace("set","")
        time_now = time_to_set.replace("alarm","")
        time_now = time_to_set.replace("for","")
        time_now = time_to_set.replace("alarm","")

        Alarm_Time = str(time_now)

        while True:
            current_time = datetime.datetime.now().strftime("%H:%M")

            if current_time == Alarm_Time:
                print("Wake Up Sir, It's Time to work.")
                PlaySound("")
                TimeHere = open('C:\\Users\\HIMI\\Desktop\\Jarvis\\data.txt','a')
                TimeHere.write(query)
                TimeHere.close()

def ChromeAuto(query):
    while True:
        query = str(query)
        if 'new tab' in query:
            press_and_release('ctrl + t')
        elif 'close tab' in query:
            press_and_release('ctrl + w')
        elif 'new window' in query:
            press_and_release('ctrl + n')
        elif 'close window' in query:
            press_and_release('ctrl + shift + w')
        elif 'history' in query:
            press_and_release('ctrl + h')
        elif 'download bar' in query:
            press_and_release('ctrl + j')
        elif 'bookmarks' in query:
            press_and_release('ctrl + d')
        elif 'favourates' in query:
            press_and_release('ctrl + shift + b')
        elif 'settings' in query:
            press_and_release('alt + e')
        elif 'find' in query:
            query = query.replace("find","")
            query = query.replace(" ","")
            press_and_release('ctrl + f')
            write(query)
        elif 'read page' in query:
            press_and_release('ctrl + shift + u')
        elif 'switch tab' in query:
            tab = query.replace("switch tab","")
            Tab = tab.replace("to","")
            Tab = Tab.replace(" ","")
            tab_num = int(Tab)
            Tab1 = f'ctrl + {tab_num}'
            press_and_release(Tab1)
        elif 'disable' in query:
            print("Auto chrome mode is disabled")
            speak("Auto chrome mode is disabled")
            pass

def YoutubeAuto(command):
    while True:
        query = str(command)
        if 'pause' in query:
            press('space bar')
        elif 'resume' in query:
            press('space bar')
        elif 'full screen' in query:
            press('f')
        elif 'film screen' in query:
            press('t')
        elif 'skip' in query:
            press('l')
        elif 'back' in query:
            press('j')
        elif 'increase' in query:
            press_and_release('shift + >')
        elif 'decrease' in query:
            press_and_release('shift + <')
        elif 'previous' in query:
            press_and_release('shift + p')
        elif 'next' in query:
            press_and_release('shift + n')
        elif 'search' in query:
            click(x=667, y=146)
            speak("what to search sir?")
            search = takeCommand()
            write(search)
            press('enter')
        elif 'mute' in query:
            press('m')
        elif 'unmute' in query:
            press('m')
        elif 'disable' in query:
            speak("YoutubeAuto mode is disabled. Now you can carry on with your work")
            pass

def My_Location():
    op = "https://www.google.com/maps/place/28%C2%B036'11.4%22N+77%C2%B003'58.7%22E/@28.6031746,77.0618113,826m/data=!3m2!1e3!4b1!4m14!1m7!3m6!1s0x390d1b3399ab2c5f:0xb829bcae933102b4!2sMahavir+Enclave+Part+3,+Mahavir+Enclave,+Delhi,+110059!3b1!8m2!3d28.6072242!4d77.0715564!3m5!1s0x0:0xfc21da43fcc58fdb!7e2!8m2!3d28.6031698!4d77.0662964"
    webbrowser.open(op)
    ip_add = requests.get('https://api.ipify.org').text
    url = 'https://get.geojs.io/v1/ip/geo/' + ip_add + '.json'
    geo_q = requests.get(url)
    geo_d = geo_q.json()
    state = geo_d['city']
    country = geo_d['country']
    print(f"Sir, you are now in {state, country}.")
    speak(f"Sir, you are now in {state, country}.")

def Google_maps(place):
    Url_place = "https://www.google.com/maps/place/" + str(place)
    geolocator = Nominatim(user_agent="myGeocoder")
    location = geolocator.geocode(place, addressdetails= True)
    target_location = location.latitude , location.longitude
    location = location.raw['address']
    target = {'city' : location.get('city',''),
                'state' : location.get('state',''),
                'country' : location.get('country','')}
    current_loca = geocoder.ip('me')
    current_location = current_loca.latlng

    distance = str(great_circle(current_location,target_location))
    distance = str(distance.split(' ',1)[0])
    distance = round(float(distance),2)

    print(target)
    speak(target)
    print(f"sir, {place} is {distance} kilometre away from your location")
    speak(f"sir, {place} is {distance} kilometre away from your location")

def OnlineClasses():
    classroom = "https://app.classroom.live/student-live-list"
    webbrowser.open(classroom)
    sleep(8)
    click(x=882, y=562)
    sleep(5)
    click(x=412, y=571)
    sleep(1)
    #click(x=491, y=575) #Optional: if you have camera
    click(x=1101, y =443)
    print("Class joined sir")
    speak("Class joined sir")

def reading_word():
    print("Which word file do you want to read.")
    speak("Which word file do you want to read.")
    Word = takeCommand()
    file_name = f"{Word}.docx"
    doc = docx.Document(file_name)
    all_paras = doc.paragraphs
    print(f"There are {len(all_paras)} paragraphs. Which para should I read")
    speak(f"There are {len(all_paras)} paragraphs. Which para should I read")
    para_num = takeCommand()
    para_num = para_num
    if 'all' in para_num:
        for para in all_paras:
            print(para.text)
            speak(para.text)
    else:
        single_para = doc.paragraphs[para_num]
        print(single_para.text)

def writing_word():
    speak("opening word file")
    speak("what should i write in file")
    write_word = takeCommand()
    mydoc = docx.Document()
    mydoc.add_paragraph(write_word)
    speak("please tell the name for file")
    path = takeCommand().lower()
    mydoc.save(path)

def zoom():
    speak("Which class do you want to join")
    class_zoom = takeCommand().lower()
    speak("please tell password")
    pass_zoom = takeCommand().lower()
    pass_zoom = pass_zoom.replace(" ","")
    click(x=751, y=450)
    sleep(1)
    click(x=711, y=392)
    if 'science' in class_zoom:
        write("857 748 6948")
        press('enter')
        sleep(3)
        click(x=802, y=391)
        sleep(0.5)
        write("1234321")
        press('enter')
    elif 'maths' in class_zoom:
        write("373 880 8637")
        press('enter')
        sleep(3)
        click(x=802, y=391)
        sleep(0.5)
        write("123456")
        press('enter')
    else:
        speak(f"soory sir, I am not able to find any class of {class_zoom}, please give Meeting id")
        class_zoom_id = takeCommand()
        write(class_zoom_id)
        press('enter')
        sleep(3)
        click(x=802, y=391)
        sleep(0.5)
        write(pass_zoom)
        press('enter')
 
def TimeTable():
    from Database.TimeTable import Time
    speak("Checking")
    value = Time()
    Noti = Notify()
    Noti.title = "Time-Table"
    Noti.message = str(value)
    Noti.send()

def game_download(query):
    query = str(query)
    query.replace(" ","-")
    webbrowser.open(f"https://oceanof-games.com/{query}/")
    press('end')
    click(x=616, y=317)
    sleep(5)
    click(x=692, y=880)
    sleep(5)
    click(x=1481, y=109)
    sleep(5)
    press('end')
    click(x=364, y=198)
    sleep(5)
    press('page down')
    sleep(1)
    press('page down')
    sleep(1)
    press('page down')
    sleep(1)
    press('page down')
    sleep(1)
    click(x=576, y=286)
    sleep(15)
    speak("Your download has started, please wait till it is downloaded.")

def whatsapp(query,contacts):
    print(f"what message do you want to send to {query}?")
    speak(f"what message do you want to send to {query}?")
    message1 = takeCommand().lower()
    webbrowser.open("https://web.whatsapp.com/")
    sleep(20)
    click(x=194, y=268)
    write(contacts[query])
    sleep(1)
    click(x=240, y=383)
    sleep(1)
    click(x=679, y=809)
    sleep(1)
    write(message1)
    sleep(2)
    press('enter')

def scroll(query):
    query = str(query)
    query = query.lower()
    if 'once' in query:
        query = query.replace('scroll', '')
        query = query.replace(' ', '')
        query = query.replace('once', '')
        if query == 'up':
            sleep(5)
            click(x=1594, y=112)
        elif query == 'down':
            sleep(5)
            click(x=1591, y=849)
    elif 'scrolling' in query:
        print("starting")
        query = query.replace('scrolling', '')
        query = query.replace('start', '')
        query = query.replace(' ', '')
        if query == 'up':
            sleep(5)
            _physically_pressed_keys('up')
        elif query == 'down':
            sleep(5)
            click(x=1591, y=849)
print("starting")
scroll("scroll up once")